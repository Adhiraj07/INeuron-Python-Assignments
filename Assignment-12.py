#!/usr/bin/env python
# coding: utf-8
1. In what modes should the PdfFileReader() and PdfFileWriter() File objects will be opened? 
Answer: read binary (rb) for PdfFileREader()
        write binary (wb) PdfFileWriter()2. From a PdfFileReader object, how do you get a Page object for page 5? 
# In[4]:


get_ipython().system('pip install pyPDF2')


# In[21]:


import PyPDF2 as pdf
pdf_obj=open('example.pdf','rb')
pdfReader=pdf.PdfReader(pdf_obj) #PDF reader

pgobj=pdfReader.pages[4]
pgobj.extract_text()


3. What PdfFileReader variable stores the number of pages in the PDF document? 
# In[25]:


pdfObj=open('example.pdf','rb')

Pdf_Reader= PyPDF2.PdfReader(pdfObj)

len(Pdf_Reader.pages)  

4. If a PdfFileReader object’s PDF is encrypted with the password swordfish, what must you do before you can obtain Page objects from it? 
Answer we can use pdfReader.decrypt("swordfish")5. What methods do you use to rotate a page? 
# In[4]:


import PyPDF2

# Open the PDF file
with open('example.pdf', 'rb') as file:
    reader = PyPDF2.PdfReader(file)
    writer = PyPDF2.PdfWriter()

    # Rotate each page and add it to the output PDF
    for page in reader.pages:
        rotated_page = page.rotate(90)  # Rotate 90 degrees clockwise
        writer.add_page(rotated_page)

    # Save the output PDF
    with open('output.pdf', 'wb') as output_file:
        writer.write(output_file)

6. What is the difference between a Run object and a Paragraph object?

In summary, a paragraph is a higher-level structure that encapsulates a block of text, while a run represents a portion of text within a paragraph with specific formatting attributes. Paragraphs organize and group text, while runs apply formatting styles to the text within a paragraph.
7. How do you obtain a list of Paragraph objects for a Document object that’s stored in a variable named doc? 
# In[15]:


get_ipython().system('pip install python-docx')


# In[24]:


import docx
doc=docx.Document('Assignment_12.docx') ## Load the document

par=doc.paragraphs

for i in par:
    print(i.text)

8. What type of object has bold, underline, italic, strike, and outline variables?

A Run object has bold, underline,italic,strike and outline variables9. What is the difference between False, True, and None for the bold variable? 

the bold variable, assigning False might imply that the text should not be displayed in bold.

the bold variable, assigning True might indicate that the text should be displayed in bold.

the bold variable, assigning None might mean that no specific value or state regarding boldness is assigned. It can be used as a placeholder value until a proper value is assigned later.

True always makes the Run object bolded and False makes it always not bolded, no matter what the style’s bold setting is. None will make the Run object just use the style’s bold setting
10. How do you create a Document object for a new Word document? 

By Calling the docx.Document() function.
11. How do you add a paragraph with the text 'Hello, there!' to a Document object stored in a variable named doc? 

# In[37]:


import docx

docum=docx.Document()

docum.add_paragraph("Hello, there!") # we are adding the object
docum.save("addingParagraph.docx") #saving the object in new doc


doc1=docx.Document('addingParagraph.docx')

para1=doc1.paragraphs

for i in para1:
    print(i.text)

12. What integers represent the levels of headings available in Word documents? 

The arguments to add_heading() are a string of the heading text and an integer from 0 to 4. The integer 0 makes the heading the Title style, which is used for the top of the document. Integers 1 to 4 are for various heading levels, with 1 being the main heading and 4 the lowest subheading
# In[42]:


from docx import Document
# Create a new document
doc = Document()

# Add a heading
heading_text = "Sample Heading"
heading_level = 2 
doc.add_heading(heading_text, level=heading_level)

# Save the document
doc.save("output.docx")


# In[ ]:




